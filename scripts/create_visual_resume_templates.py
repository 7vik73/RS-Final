"""Create 25 visually distinct DOCX resume templates for ResumeIQ.

The documents are intentionally placeholder templates, not fake resumes. They
are useful for testing DOCX parsing after a user fills in real candidate data.
"""

from pathlib import Path
from zipfile import ZipFile, ZIP_DEFLATED

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


OUTPUT_DIR = Path(__file__).resolve().parent.parent / "visual_resume_templates"
ZIP_PATH = Path(__file__).resolve().parent.parent / "visual_resume_templates.zip"


DOMAINS = [
    ("01_backend_python_developer", "Backend Python Developer", ["python", "flask", "fastapi", "sqlite", "postgresql", "rest api", "git"]),
    ("02_frontend_react_developer", "Frontend React Developer", ["javascript", "typescript", "react", "tailwind", "html", "css", "git"]),
    ("03_full_stack_engineer", "Full Stack Engineer", ["python", "javascript", "flask", "react", "sqlite", "api", "git"]),
    ("04_data_analyst", "Data Analyst", ["python", "sql", "excel", "power bi", "tableau", "pandas", "matplotlib"]),
    ("05_data_scientist", "Data Scientist", ["python", "machine learning", "nlp", "pandas", "numpy", "scikit-learn", "matplotlib"]),
    ("06_machine_learning_engineer", "Machine Learning Engineer", ["python", "machine learning", "deep learning", "pytorch", "tensorflow", "scikit-learn", "docker"]),
    ("07_nlp_engineer", "NLP Engineer", ["python", "nlp", "nltk", "transformers", "sentence transformers", "scikit-learn", "api"]),
    ("08_devops_engineer", "DevOps Engineer", ["linux", "git", "docker", "kubernetes", "aws", "ci/cd", "monitoring"]),
    ("09_cloud_engineer", "Cloud Engineer", ["aws", "azure", "gcp", "linux", "docker", "terraform", "networking"]),
    ("10_cybersecurity_analyst", "Cybersecurity Analyst", ["linux", "network security", "siem", "incident response", "python", "risk assessment", "firewalls"]),
    ("11_database_administrator", "Database Administrator", ["sql", "mysql", "postgresql", "oracle", "backup", "performance tuning", "linux"]),
    ("12_mobile_app_developer", "Mobile App Developer", ["kotlin", "swift", "flutter", "firebase", "api integration", "git", "mobile ui"]),
    ("13_ui_ux_designer", "UI/UX Designer", ["figma", "wireframes", "prototyping", "user research", "design systems", "usability testing", "accessibility"]),
    ("14_graphic_designer", "Graphic Designer", ["photoshop", "illustrator", "branding", "layout design", "typography", "figma", "creative assets"]),
    ("15_digital_marketing_specialist", "Digital Marketing Specialist", ["seo", "google analytics", "content marketing", "social media", "email marketing", "campaigns", "excel"]),
    ("16_hr_recruiter", "HR Recruiter", ["recruitment", "screening", "interviewing", "onboarding", "ats", "communication", "excel"]),
    ("17_business_analyst", "Business Analyst", ["requirements gathering", "stakeholder management", "sql", "excel", "jira", "process mapping", "documentation"]),
    ("18_project_manager", "Project Manager", ["agile", "scrum", "jira", "risk management", "planning", "stakeholder communication", "reporting"]),
    ("19_product_manager", "Product Manager", ["roadmap", "user research", "analytics", "prioritization", "agile", "wireframes", "stakeholders"]),
    ("20_finance_analyst", "Finance Analyst", ["excel", "financial modeling", "forecasting", "budgeting", "variance analysis", "power bi", "reporting"]),
    ("21_sales_executive", "Sales Executive", ["lead generation", "crm", "negotiation", "pipeline management", "communication", "market research", "reporting"]),
    ("22_customer_support_specialist", "Customer Support Specialist", ["customer service", "ticketing", "communication", "crm", "troubleshooting", "documentation", "sla"]),
    ("23_content_writer", "Content Writer", ["copywriting", "seo", "editing", "research", "blog writing", "content strategy", "wordpress"]),
    ("24_quality_assurance_tester", "Quality Assurance Tester", ["manual testing", "automation testing", "selenium", "api testing", "bug reporting", "jira", "test cases"]),
    ("25_operations_manager", "Operations Manager", ["process improvement", "vendor management", "excel", "reporting", "team management", "inventory", "planning"]),
]


PALETTES = [
    ("0F766E", "CCFBF1", "134E4A"),
    ("1D4ED8", "DBEAFE", "1E3A8A"),
    ("7C3AED", "EDE9FE", "4C1D95"),
    ("C2410C", "FFEDD5", "7C2D12"),
    ("047857", "D1FAE5", "064E3B"),
    ("BE123C", "FFE4E6", "881337"),
    ("4338CA", "E0E7FF", "312E81"),
    ("334155", "E2E8F0", "0F172A"),
    ("A16207", "FEF3C7", "713F12"),
    ("0E7490", "CFFAFE", "164E63"),
]


LAYOUTS = [
    "left_sidebar",
    "right_sidebar",
    "top_band",
    "split_header",
    "minimal_line",
]


def shade_cell(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shading = OxmlElement("w:shd")
    shading.set(qn("w:fill"), fill)
    tc_pr.append(shading)


def set_cell_border(cell, color="FFFFFF"):
    tc_pr = cell._tc.get_or_add_tcPr()
    borders = OxmlElement("w:tcBorders")
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        tag = OxmlElement(f"w:{edge}")
        tag.set(qn("w:val"), "single")
        tag.set(qn("w:sz"), "4")
        tag.set(qn("w:space"), "0")
        tag.set(qn("w:color"), color)
        borders.append(tag)
    tc_pr.append(borders)


def set_table_width(table, widths):
    for row in table.rows:
        for index, width in enumerate(widths):
            row.cells[index].width = Inches(width)


def add_run(paragraph, text, bold=False, size=10, color="000000"):
    run = paragraph.add_run(text)
    run.bold = bold
    run.font.name = "Arial"
    run.font.size = Pt(size)
    run.font.color.rgb = RGBColor.from_string(color)
    return run


def add_heading(paragraph_or_cell, text, accent):
    paragraph = paragraph_or_cell.add_paragraph()
    paragraph.paragraph_format.space_before = Pt(8)
    paragraph.paragraph_format.space_after = Pt(3)
    add_run(paragraph, text.upper(), bold=True, size=10, color=accent)
    return paragraph


def add_body(paragraph_or_cell, text, size=9):
    paragraph = paragraph_or_cell.add_paragraph()
    paragraph.paragraph_format.space_after = Pt(3)
    add_run(paragraph, text, size=size, color="111827")
    return paragraph


def add_bullet(paragraph_or_cell, text, size=9):
    paragraph = paragraph_or_cell.add_paragraph(style="List Bullet")
    paragraph.paragraph_format.space_after = Pt(2)
    if paragraph.runs:
        paragraph.runs[0].text = ""
    add_run(paragraph, text, size=size, color="111827")
    return paragraph


def setup_doc(document):
    section = document.sections[0]
    section.top_margin = Inches(0.55)
    section.bottom_margin = Inches(0.55)
    section.left_margin = Inches(0.55)
    section.right_margin = Inches(0.55)

    normal = document.styles["Normal"]
    normal.font.name = "Arial"
    normal.font.size = Pt(9)
    normal.paragraph_format.space_after = Pt(3)
    normal.paragraph_format.line_spacing = 1.05


def fill_sidebar(cell, title, skills, accent, light, dark):
    shade_cell(cell, light)
    set_cell_border(cell, light)
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.TOP

    add_heading(cell, "Contact", dark)
    add_body(cell, "[your.email@example.com]")
    add_body(cell, "[+91 00000 00000]")
    add_body(cell, "[City, Country]")
    add_body(cell, "[LinkedIn/GitHub/Portfolio]")

    add_heading(cell, "Core Skills", dark)
    for skill in skills:
        add_bullet(cell, skill)

    add_heading(cell, "Tools", dark)
    add_body(cell, "[Add tools, platforms, IDEs, or domain software]")

    add_heading(cell, "Certifications", dark)
    add_body(cell, "[Certification Name]")
    add_body(cell, "[Certification Name]")


def fill_main(cell, title, skills, accent, style_note):
    set_cell_border(cell, "FFFFFF")
    name = cell.add_paragraph()
    name.paragraph_format.space_after = Pt(0)
    add_run(name, "[Your Full Name]", bold=True, size=22, color=accent)

    role = cell.add_paragraph()
    role.paragraph_format.space_after = Pt(8)
    add_run(role, title, bold=True, size=12, color="334155")

    add_heading(cell, "Professional Summary", accent)
    add_body(cell, f"[Write 3-4 lines positioning yourself as a {title}. Mention measurable impact, domain strengths, and relevant keywords.]")

    add_heading(cell, "Experience", accent)
    add_body(cell, f"[{title} / Internship / Relevant Role] | [Organization] | [Month Year - Month Year]", size=9)
    add_bullet(cell, "[Describe one responsibility using action verbs and role-specific tools.]")
    add_bullet(cell, "[Describe one measurable achievement or outcome.]")
    add_bullet(cell, "[Mention teamwork, ownership, or process improvement.]")

    add_heading(cell, "Projects", accent)
    add_body(cell, f"[{title} Portfolio Project]", size=9)
    add_bullet(cell, f"[Describe a project using keywords like {', '.join(skills[:4])}.]")
    add_bullet(cell, "[Mention the problem solved, your contribution, and result.]")

    add_heading(cell, "Education", accent)
    add_body(cell, "[Degree / Course] | [Institution Name] | [Graduation Year]")

    add_heading(cell, "Template Note", accent)
    add_body(cell, style_note, size=8)


def make_left_sidebar(document, title, skills, accent, light, dark):
    table = document.add_table(rows=1, cols=2)
    table.autofit = False
    set_table_width(table, [2.05, 5.35])
    fill_sidebar(table.cell(0, 0), title, skills, accent, light, dark)
    fill_main(table.cell(0, 1), title, skills, accent, "Visual template with a left skill sidebar.")


def make_right_sidebar(document, title, skills, accent, light, dark):
    table = document.add_table(rows=1, cols=2)
    table.autofit = False
    set_table_width(table, [5.35, 2.05])
    fill_main(table.cell(0, 0), title, skills, accent, "Visual template with a right contact and skill panel.")
    fill_sidebar(table.cell(0, 1), title, skills, accent, light, dark)


def make_top_band(document, title, skills, accent, light, dark):
    header = document.add_table(rows=1, cols=1)
    shade_cell(header.cell(0, 0), accent)
    set_cell_border(header.cell(0, 0), accent)
    p = header.cell(0, 0).add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    add_run(p, "[Your Full Name]", bold=True, size=22, color="FFFFFF")
    p2 = header.cell(0, 0).add_paragraph()
    p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    add_run(p2, f"{title} | [your.email@example.com] | [+91 00000 00000]", bold=True, size=10, color="FFFFFF")

    body = document.add_table(rows=1, cols=2)
    body.autofit = False
    set_table_width(body, [4.75, 2.65])
    fill_main(body.cell(0, 0), title, skills, accent, "Visual template with a bold top identity band.")
    fill_sidebar(body.cell(0, 1), title, skills, accent, light, dark)


def make_split_header(document, title, skills, accent, light, dark):
    head = document.add_table(rows=1, cols=2)
    head.autofit = False
    set_table_width(head, [4.3, 3.1])
    shade_cell(head.cell(0, 0), accent)
    shade_cell(head.cell(0, 1), light)
    set_cell_border(head.cell(0, 0), accent)
    set_cell_border(head.cell(0, 1), light)
    add_run(head.cell(0, 0).add_paragraph(), "[Your Full Name]", bold=True, size=20, color="FFFFFF")
    add_run(head.cell(0, 0).add_paragraph(), title, bold=True, size=11, color="FFFFFF")
    add_body(head.cell(0, 1), "[your.email@example.com] | [+91 00000 00000]")
    add_body(head.cell(0, 1), "[City, Country] | [LinkedIn/GitHub/Portfolio]")

    body = document.add_table(rows=1, cols=2)
    body.autofit = False
    set_table_width(body, [2.45, 4.95])
    fill_sidebar(body.cell(0, 0), title, skills, accent, "FFFFFF", dark)
    fill_main(body.cell(0, 1), title, skills, accent, "Visual template with a split header and compact skill column.")


def make_minimal_line(document, title, skills, accent, light, dark):
    p = document.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    add_run(p, "[Your Full Name]", bold=True, size=22, color="111827")
    role = document.add_paragraph()
    role.alignment = WD_ALIGN_PARAGRAPH.CENTER
    add_run(role, f"{title} | [your.email@example.com] | [+91 00000 00000]", size=10, color="475569")

    line = document.add_table(rows=1, cols=1)
    shade_cell(line.cell(0, 0), accent)
    set_cell_border(line.cell(0, 0), accent)
    line.cell(0, 0).height = Inches(0.05)

    body = document.add_table(rows=1, cols=2)
    body.autofit = False
    set_table_width(body, [5.0, 2.4])
    fill_main(body.cell(0, 0), title, skills, accent, "Minimal visual template with a clean accent divider.")
    fill_sidebar(body.cell(0, 1), title, skills, accent, light, dark)


LAYOUT_BUILDERS = {
    "left_sidebar": make_left_sidebar,
    "right_sidebar": make_right_sidebar,
    "top_band": make_top_band,
    "split_header": make_split_header,
    "minimal_line": make_minimal_line,
}


def create_visual_template(index, filename, title, skills):
    document = Document()
    setup_doc(document)
    accent, light, dark = PALETTES[index % len(PALETTES)]
    layout = LAYOUTS[index % len(LAYOUTS)]

    LAYOUT_BUILDERS[layout](document, title, skills, accent, light, dark)

    document.core_properties.title = f"{title} Visual Resume Template"
    document.core_properties.subject = "ResumeIQ visual placeholder resume template"
    document.core_properties.author = "ResumeIQ"
    document.save(OUTPUT_DIR / f"{filename}_visual.docx")


def main():
    OUTPUT_DIR.mkdir(parents=True, exist_ok=True)
    for existing in OUTPUT_DIR.glob("*.docx"):
        existing.unlink()

    for index, (filename, title, skills) in enumerate(DOMAINS):
        create_visual_template(index, filename, title, skills)

    if ZIP_PATH.exists():
        ZIP_PATH.unlink()
    with ZipFile(ZIP_PATH, "w", ZIP_DEFLATED) as zip_file:
        for path in sorted(OUTPUT_DIR.glob("*.docx")):
            zip_file.write(path, arcname=path.name)

    print(f"Created {len(list(OUTPUT_DIR.glob('*.docx')))} visual DOCX templates")
    print(ZIP_PATH)


if __name__ == "__main__":
    main()
