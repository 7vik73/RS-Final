"""Create clearly labeled DOCX resume templates for ResumeIQ testing.

These files are intentionally templates, not fabricated resumes for fake
people. They use placeholders so users can add real candidate details before
running analytics.
"""

from pathlib import Path
from zipfile import ZipFile, ZIP_DEFLATED

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Pt, RGBColor


OUTPUT_DIR = Path(__file__).resolve().parent.parent / "sample_resume_templates"
ZIP_PATH = Path(__file__).resolve().parent.parent / "sample_resume_templates.zip"


DOMAINS = [
    ("01_backend_python_developer", "Backend Python Developer", ["python", "flask", "fastapi", "sqlite", "postgresql", "rest api", "git"]),
    ("02_frontend_react_developer", "Frontend React Developer", ["javascript", "typescript", "react", "tailwind", "html", "css", "git"]),
    ("03_full_stack_engineer", "Full Stack Engineer", ["python", "javascript", "flask", "react", "sqlite", "api", "git"]),
    ("04_data_analyst", "Data Analyst", ["python", "sql", "excel", "power bi", "tableau", "pandas", "matplotlib"]),
    ("05_data_scientist", "Data Scientist", ["python", "machine learning", "nlp", "pandas", "numpy", "scikit-learn", "matplotlib"]),
    ("06_machine_learning_engineer", "Machine Learning Engineer", ["python", "machine learning", "deep learning", "pytorch", "tensorflow", "scikit-learn", "docker"]),
    ("07_nlp_engineer", "NLP Engineer", ["python", "nlp", "nltk", "transformers", "sentence transformers", "scikit-learn", "api"]),
    ("08_devops_engineer", "DevOps Engineer", ["linux", "git", "docker", "kubernetes", "aws", "ci/cd", "monitoring"]),
    ("09_cloud_engineer", "Cloud Engineer", ["aws", "azure", "gcp", "linux", "docker", "terraform", "networking"]),
    ("10_cybersecurity_analyst", "Cybersecurity Analyst", ["linux", "network security", "siem", "incident response", "python", "risk assessment", "firewalls"]),
    ("11_database_administrator", "Database Administrator", ["sql", "mysql", "postgresql", "oracle", "backup", "performance tuning", "linux"]),
    ("12_mobile_app_developer", "Mobile App Developer", ["kotlin", "swift", "flutter", "firebase", "api integration", "git", "mobile ui"]),
    ("13_ui_ux_designer", "UI/UX Designer", ["figma", "wireframes", "prototyping", "user research", "design systems", "usability testing", "accessibility"]),
    ("14_graphic_designer", "Graphic Designer", ["photoshop", "illustrator", "branding", "layout design", "typography", "figma", "creative assets"]),
    ("15_digital_marketing_specialist", "Digital Marketing Specialist", ["seo", "google analytics", "content marketing", "social media", "email marketing", "campaigns", "excel"]),
    ("16_hr_recruiter", "HR Recruiter", ["recruitment", "screening", "interviewing", "onboarding", "ats", "communication", "excel"]),
    ("17_business_analyst", "Business Analyst", ["requirements gathering", "stakeholder management", "sql", "excel", "jira", "process mapping", "documentation"]),
    ("18_project_manager", "Project Manager", ["agile", "scrum", "jira", "risk management", "planning", "stakeholder communication", "reporting"]),
    ("19_product_manager", "Product Manager", ["roadmap", "user research", "analytics", "prioritization", "agile", "wireframes", "stakeholders"]),
    ("20_finance_analyst", "Finance Analyst", ["excel", "financial modeling", "forecasting", "budgeting", "variance analysis", "power bi", "reporting"]),
    ("21_sales_executive", "Sales Executive", ["lead generation", "crm", "negotiation", "pipeline management", "communication", "market research", "reporting"]),
    ("22_customer_support_specialist", "Customer Support Specialist", ["customer service", "ticketing", "communication", "crm", "troubleshooting", "documentation", "sla"]),
    ("23_content_writer", "Content Writer", ["copywriting", "seo", "editing", "research", "blog writing", "content strategy", "wordpress"]),
    ("24_quality_assurance_tester", "Quality Assurance Tester", ["manual testing", "automation testing", "selenium", "api testing", "bug reporting", "jira", "test cases"]),
    ("25_operations_manager", "Operations Manager", ["process improvement", "vendor management", "excel", "reporting", "team management", "inventory", "planning"]),
]


def set_document_styles(document):
    """Apply a simple Google Docs style preset."""
    section = document.sections[0]
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)

    normal = document.styles["Normal"]
    normal.font.name = "Arial"
    normal.font.size = Pt(11)
    normal.paragraph_format.space_after = Pt(8)
    normal.paragraph_format.line_spacing = 1.15

    for style_name, size in [("Heading 1", 16), ("Heading 2", 14), ("Heading 3", 12)]:
        style = document.styles[style_name]
        style.font.name = "Arial"
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = RGBColor(0, 0, 0)


def add_placeholder_line(document, label, value):
    paragraph = document.add_paragraph()
    run = paragraph.add_run(f"{label}: ")
    run.bold = True
    paragraph.add_run(value)


def add_bullets(document, items):
    for item in items:
        document.add_paragraph(item, style="List Bullet")


def create_template(filename, title, skills):
    document = Document()
    set_document_styles(document)

    heading = document.add_paragraph()
    heading.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = heading.add_run("[Your Full Name]")
    run.bold = True
    run.font.size = Pt(20)

    subtitle = document.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle.add_run(title).bold = True

    contact = document.add_paragraph()
    contact.alignment = WD_ALIGN_PARAGRAPH.CENTER
    contact.add_run("[your.email@example.com] | [+91 00000 00000] | [City, Country] | [LinkedIn/GitHub/Portfolio]")

    document.add_heading("Professional Summary", level=1)
    document.add_paragraph(
        f"[Write 3-4 lines describing your background as a {title}. Mention your strongest domain skills, work style, and measurable impact.]"
    )

    document.add_heading("Skills", level=1)
    add_bullets(document, [", ".join(skills), "[Add more tools, platforms, or domain keywords relevant to your experience.]"])

    document.add_heading("Experience", level=1)
    add_placeholder_line(document, "Role", f"[{title} / Internship / Relevant Role]")
    add_placeholder_line(document, "Organization", "[Company or College Organization]")
    add_placeholder_line(document, "Duration", "[Month Year - Month Year]")
    add_bullets(
        document,
        [
            "[Describe one responsibility using action verbs and relevant tools.]",
            "[Describe one measurable achievement or outcome.]",
            "[Mention collaboration, ownership, or process improvement if applicable.]",
        ],
    )

    document.add_heading("Projects", level=1)
    add_placeholder_line(document, "Project", f"[{title} Portfolio Project]")
    add_bullets(
        document,
        [
            f"[Describe a project related to {title.lower()} using keywords such as {', '.join(skills[:4])}.]",
            "[Mention the problem solved, your contribution, and the result.]",
        ],
    )

    document.add_heading("Education", level=1)
    add_placeholder_line(document, "Degree", "[Degree / Certification / Course]")
    add_placeholder_line(document, "Institution", "[Institution Name]")
    add_placeholder_line(document, "Year", "[Graduation Year]")

    document.add_heading("Certifications", level=1)
    add_bullets(document, ["[Certification Name - Issuing Organization]", "[Certification Name - Issuing Organization]"])

    document.add_heading("Additional Details", level=1)
    add_bullets(document, ["[Languages]", "[Awards or achievements]", "[Availability or preferred location]"])

    document.core_properties.title = f"{title} Resume Template"
    document.core_properties.subject = "ResumeIQ placeholder resume template"
    document.core_properties.author = "ResumeIQ"

    path = OUTPUT_DIR / f"{filename}.docx"
    document.save(path)
    return path


def main():
    OUTPUT_DIR.mkdir(parents=True, exist_ok=True)
    for existing in OUTPUT_DIR.glob("*.docx"):
        existing.unlink()

    created = [create_template(*domain) for domain in DOMAINS]

    if ZIP_PATH.exists():
        ZIP_PATH.unlink()
    with ZipFile(ZIP_PATH, "w", ZIP_DEFLATED) as zip_file:
        for path in created:
            zip_file.write(path, arcname=path.name)

    print(f"Created {len(created)} DOCX templates")
    print(ZIP_PATH)


if __name__ == "__main__":
    main()
