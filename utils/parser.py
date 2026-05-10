"""Resume parsing helpers for PDF and DOCX files."""

import json
import re
from pathlib import Path

import pdfplumber
from docx import Document

from utils.skills import extract_skills


SECTION_PATTERNS = {
    "experience": r"(experience|work history|employment)(.*?)(projects|education|skills|certifications|$)",
    "projects": r"(projects|academic projects)(.*?)(experience|education|skills|certifications|$)",
    "education": r"(education|academics)(.*?)(experience|projects|skills|certifications|$)",
    "certifications": r"(certifications|certificates)(.*?)(experience|projects|education|skills|$)",
}


def parse_pdf(path):
    """Extract text from every page of a PDF resume."""
    text_parts = []
    with pdfplumber.open(path) as pdf:
        for page in pdf.pages:
            text_parts.append(page.extract_text() or "")
    return "\n".join(text_parts).strip()


def parse_docx(path):
    """Extract paragraph and table text from a DOCX resume."""
    document = Document(path)
    parts = [paragraph.text for paragraph in document.paragraphs]

    for table in document.tables:
        for row in table.rows:
            for cell in row.cells:
                parts.append(cell.text)

    return "\n".join(parts).strip()


def parse_resume(path):
    """Parse a supported resume file and return structured information."""
    suffix = Path(path).suffix.lower()
    if suffix == ".pdf":
        raw_text = parse_pdf(path)
    elif suffix == ".docx":
        raw_text = parse_docx(path)
    else:
        raise ValueError("Only PDF and DOCX resumes are supported.")

    return {
        "raw_text": raw_text,
        "email": extract_email(raw_text),
        "phone": extract_phone(raw_text),
        "skills": extract_skills(raw_text),
        "sections": extract_sections(raw_text),
    }


def extract_email(text):
    match = re.search(r"[\w.+-]+@[\w-]+\.[\w.-]+", text)
    return match.group(0) if match else ""


def extract_phone(text):
    match = re.search(r"(\+?\d[\d\s().-]{8,}\d)", text)
    return match.group(0).strip() if match else ""


def extract_candidate_name(filename, raw_text):
    """Prefer the first useful resume line, then fall back to the filename."""
    for line in raw_text.splitlines()[:5]:
        cleaned = line.strip()
        if 2 <= len(cleaned.split()) <= 4 and "@" not in cleaned:
            return cleaned[:80]
    return Path(filename).stem.replace("_", " ").replace("-", " ").title()


def extract_sections(text):
    """Extract common resume sections with readable regex rules."""
    normalized = re.sub(r"\s+", " ", text.lower())
    sections = {}

    for section_name, pattern in SECTION_PATTERNS.items():
        match = re.search(pattern, normalized, re.IGNORECASE)
        sections[section_name] = match.group(2).strip() if match else ""

    return sections


def sections_to_json(sections):
    return json.dumps(sections)
